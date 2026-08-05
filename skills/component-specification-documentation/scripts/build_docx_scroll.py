"""
Build the component-specification .docx so it mirrors the Confluence/Scroll export format that TM Forum
actually publishes for these components (reference:
`TMFC003 Product Order Delivery Orchestration and Management v3.0.0-v9-*.docx`).

Why this exists alongside build_docx.py: build_docx.py declares its own styling through docx-js, which
cannot reproduce the Scroll export's look -- that format's identity lives in its *styles*
(`Title`, `Subline Header`, `SublineHeader Level2`, auto-numbered `Heading 1-3`, `Scroll Table Normal`,
`Scroll Panel`) and in the numbering definitions those Heading styles point at. Re-declaring all of that
by hand is both laborious and fragile. So this script uses the reference document itself as a **template**:
open it with python-docx, strip its body content, and append ours using its styles. Everything visual then
matches by construction.

Key consequences of templating:
  - **Headings carry no literal section number.** The template's Heading 1-3 styles auto-number via
    `<w:numPr>`, so emitting "1. Overview" would render as "1. 1. Overview". The leading number is
    stripped from every heading; Word regenerates it. (The PDF still needs literal numbers, which is why
    the shared .md keeps them and only this script strips them.)
  - **No cover table and no Notice page**, per user decision: the PDF is the published artifact and keeps
    that furniture, while the .docx mirrors the Scroll export, which has neither.
  - The reference's "Exported on <timestamp>" line is deliberately **not** reproduced. It is an export
    artifact rather than document content, and a changing timestamp would make every regeneration show as
    a diff.

The two source .md files are never modified, same as build_pdf.py and build_docx.py.

Usage:
    python build_docx_scroll.py path/to/TMFCxxx_Name.md [template.docx] [output.docx]
"""
import copy
import os
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from build_docx import _parse_blocks
from build_pdf import (_find_component_yaml, _load_component_metadata, _spaced_name,
                       _strip_front_matter)

DEFAULT_TEMPLATE = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "..", "assets",
    "scroll_export_template.docx")

TITLE_STYLE = "Title"
SUBLINE_STYLE = "Subline Header"
TABLE_STYLE = "Scroll Table Normal"
PANEL_STYLE = "Scroll Panel"
IMAGE_WIDTH_IN = 5.90          # every image in the reference export is exactly this wide
SUBLINE_TEXT = "Components and Canvas"
PANEL_TITLE = "In this deliverable:"

_LEADING_NUMBER = re.compile(r"^\d+(?:\.\d+)*\.?\s+")
_SOURCE_CAPTION = re.compile(r"^\s*(PlantUML|SVG|Mermaid)\s+source\s*:", re.IGNORECASE)


def _strip_section_number(text):
    """'3.4.1. Published Events' -> 'Published Events'. The template's Heading styles auto-number, so a
    literal number here would double up."""
    return _LEADING_NUMBER.sub("", text, count=1)


def _clear_body(doc):
    """Remove every paragraph and table from the body, leaving the trailing <w:sectPr> (and therefore the
    page setup and the header/footer references) intact.

    Note this also removes the paragraph that carries the reference's *first* section break, collapsing
    the document to a single section -- which is what we want: one header/footer for the whole document
    rather than the export's two."""
    body = doc.element.body
    for child in list(body.iterchildren()):
        # w:sdt matters as much as w:p/w:tbl: Word wraps a table-of-contents field in a structured
        # document tag, so clearing only paragraphs and tables leaves the template's original ToC
        # sitting at the top of the generated document.
        if child.tag in (qn("w:p"), qn("w:tbl"), qn("w:sdt")):
            body.remove(child)


def _exclude_from_toc(paragraph):
    """Force outline level to 'body text' (9).

    The template's Title and Subline Header styles carry an outline level, so a `TOC \\o "1-3"` field
    picks them up and the contents list opens with the document title and subtitle as if they were
    sections. Overriding the level on the paragraph itself keeps the styles' appearance while taking
    them out of the ToC."""
    p_pr = paragraph._p.get_or_add_pPr()
    for existing in p_pr.findall(qn("w:outlineLvl")):
        p_pr.remove(existing)
    lvl = p_pr.makeelement(qn("w:outlineLvl"), {qn("w:val"): "9"})
    p_pr.append(lvl)


def _set_text_preserving_style(paragraph, text):
    """Replace a paragraph's text but keep its first run's formatting."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for extra in runs[1:]:
        extra._element.getparent().remove(extra._element)


def _retitle_header_footer(doc, title_line):
    """Set the running header to this document's own title.

    Two things to handle. The template's header names the version it was exported at, so it must be
    rewritten. And collapsing the export's two sections down to one leaves the *second* section's
    sectPr, whose header is empty -- so this can't just edit existing text, it has to create the
    paragraph when there isn't one. The footer is left alone: the template's is a STYLEREF + page-number
    field pair ("Administrative Appendix - 26") that keeps working as-is."""
    for section in doc.sections:
        header = section.header
        if header is None:
            continue
        header.is_linked_to_previous = False
        target = next((p for p in header.paragraphs if p.text.strip()), None)
        if target is not None:
            _set_text_preserving_style(target, title_line)
        else:
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.add_run(title_line)


def _add_runs(paragraph, runs):
    for r in runs:
        if r.get("break"):
            # A same-paragraph line break (build_docx.py's _parse_inline turns <br> into this marker),
            # for multi-value cells (resources/operations/event names, one entry per line -- see
            # SKILL.md) -- not a new paragraph, so it stays one row/one cell, just multiple lines.
            paragraph.add_run().add_break()
            continue
        run = paragraph.add_run(r.get("text", ""))
        if r.get("bold"):
            run.bold = True
        if r.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def _add_toc(doc):
    """A real Word TableOfContents field, inside the reference's 'In this deliverable:' panel. Word
    computes and refreshes the page numbers itself, so there is no page-location pass to do here (the
    PDF build needs one only because xhtml2pdf has no equivalent)."""
    try:
        table = doc.add_table(rows=2, cols=1, style=PANEL_STYLE)
    except KeyError:
        table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_or_add(table.cell(0, 0), PANEL_TITLE, bold=True)

    cell = table.cell(1, 0)
    p = cell.paragraphs[0]
    run = p.add_run()
    fld = run._element.makeelement(qn("w:fldSimple"), {})
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    inner_p = fld.makeelement(qn("w:p"), {})
    inner_r = inner_p.makeelement(qn("w:r"), {})
    inner_t = inner_r.makeelement(qn("w:t"), {})
    inner_t.text = "Right-click and choose Update Field to build the table of contents."
    inner_r.append(inner_t)
    inner_p.append(inner_r)
    fld.append(inner_p)
    run._element.addnext(fld)


def _set_table_look(table):
    """Enable only the header-row conditional format.

    `Scroll Table Normal` also defines a *first column* format (bold, coloured), which Word applies by
    default and which makes the leading data cell of every table look like a heading -- the component
    name in the Overview table rendered bold blue. Only firstRow is wanted here."""
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblLook")):
        tbl_pr.remove(existing)
    look = tbl_pr.makeelement(qn("w:tblLook"), {
        qn("w:val"): "04A0", qn("w:firstRow"): "1", qn("w:lastRow"): "0",
        qn("w:firstColumn"): "0", qn("w:lastColumn"): "0",
        qn("w:noHBand"): "0", qn("w:noVBand"): "1",
    })
    tbl_pr.append(look)


def _set_or_add(cell, text, bold=False):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold


def _set_single_line_borders(table):
    """Force every table to render with a single 1px line on every edge (outer + inside), matching
    build_pdf.py's `border-collapse: collapse` + `border: 1px solid #999999` rule -- don't rely on
    `Scroll Table Normal`'s own border definition, since a template swap or style edit could silently
    drop or double it (e.g. a `double`/`thick` val, or no `insideH`/`insideV`, leaving ruled rows only)."""
    tbl_pr = table._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single", qn("w:sz"): "4", qn("w:space"): "0", qn("w:color"): "999999",
        })
        borders.append(el)
    tbl_pr.append(borders)


def _add_table(doc, block):
    header, rows = block["header"], block["rows"]
    n_cols = len(header)
    try:
        table = doc.add_table(rows=1, cols=n_cols, style=TABLE_STYLE)
    except KeyError:
        table = doc.add_table(rows=1, cols=n_cols)
    table.autofit = True
    _set_table_look(table)
    _set_single_line_borders(table)

    for j, cell_runs in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        _add_runs(cell.paragraphs[0], [{**r, "bold": True} for r in cell_runs])
    # repeat the header row when a long table breaks across pages
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(tr_pr.makeelement(qn("w:tblHeader"), {}))

    widths = block.get("colWidthsPct") or []
    for row_runs in rows:
        cells = table.add_row().cells
        for j in range(n_cols):
            cell = cells[j]
            cell.paragraphs[0].text = ""
            if j < len(row_runs):
                _add_runs(cell.paragraphs[0], row_runs[j])
    if widths and len(widths) == n_cols:
        total = Inches(6.5)
        for j, pct in enumerate(widths):
            w = int(total * pct / 100)
            for row in table.rows:
                row.cells[j].width = w
    return table


def build_docx_scroll(md_path, template_path=None, out_path=None, yaml_path=None,
                      supplement_path=None):
    md_path = os.path.abspath(md_path)
    base_dir = os.path.dirname(md_path)  # Diagrams/temp/ -- see build_pdf.reset_temp_dir()
    diagrams_dir = os.path.dirname(base_dir)
    stem = os.path.splitext(os.path.basename(md_path))[0]

    if supplement_path is None:
        # Supplement.md is hand-maintained and stays directly under Diagrams/, one level above the
        # .md's own Diagrams/temp/.
        supplement_path = os.path.join(diagrams_dir, f"{stem}_Supplement.md")
    if not os.path.exists(supplement_path):
        raise FileNotFoundError(
            f"Supplement file not found: {supplement_path}\n"
            "Chapters 5.2/5.3/6 live there; seed it from templates/Supplement_Template.md.")
    if yaml_path is None:
        # takes the .md's directory, then looks two levels up: the component root holds the single
        # componentMetadata YAML, while base_dir is full of unrelated diagram-source .yaml files and
        # the level directly above it (Diagrams/) holds the .docx and hand-maintained files, not the
        # component YAML.
        yaml_path = _find_component_yaml(base_dir)
    if template_path is None:
        template_path = DEFAULT_TEMPLATE
    template_path = os.path.abspath(template_path)
    if not os.path.exists(template_path):
        raise FileNotFoundError(
            f"Scroll-export template not found: {template_path}\n"
            "This format is produced by templating from a real Confluence/Scroll export -- place one at "
            "that path (see this module's docstring).")
    if out_path is None:
        # The .docx has no interim stop in temp/ -- per explicit user instruction it's a final
        # deliverable that stays directly under Diagrams/, same level as the hand-maintained files.
        out_path = os.path.join(diagrams_dir, f"{stem}.docx")

    meta = _load_component_metadata(yaml_path)
    display = _spaced_name(meta["name"])
    version = str(meta.get("version", "")).strip()
    title_line = f"{meta['id']} {display} v{version}" if version else f"{meta['id']} {display}"

    with open(md_path, encoding="utf-8") as f:
        main_md = _strip_front_matter(f.read())
    with open(supplement_path, encoding="utf-8") as f:
        supp_md = _strip_front_matter(f.read())
    blocks = _parse_blocks(main_md.rstrip() + "\n\n" + supp_md.lstrip(), base_dir)

    doc = Document(template_path)
    _clear_body(doc)
    _retitle_header_footer(doc, f"{SUBLINE_TEXT} – {title_line}")

    def para(text, style=None):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        if text:
            p.add_run(text)
        return p

    try:
        title_p = para(title_line, TITLE_STYLE)
    except KeyError:
        title_p = para(title_line)
    try:
        subline_p = para(SUBLINE_TEXT, SUBLINE_STYLE)
    except KeyError:
        subline_p = para(SUBLINE_TEXT)
    _exclude_from_toc(title_p)
    _exclude_from_toc(subline_p)
    _add_toc(doc)

    skipped_images = []
    for block in blocks:
        kind = block["type"]
        if kind == "heading":
            # the document's own H1 title is emitted above via the Title style; a markdown H1 here is
            # that same title and would otherwise become a spurious top-level ToC entry
            if block["level"] == 1:
                continue
            # shift down one: markdown ## is a top-level numbered section, which the reference export
            # styles as Heading 1 (its "Overview" is Heading 1, not Heading 2)
            level = min(block["level"] - 1, 3)
            text = _strip_section_number("".join(r.get("text", "") for r in block["runs"]))
            p = doc.add_paragraph(style=f"Heading {level}")
            p.add_run(text)
        elif kind == "para":
            text = "".join(r.get("text", "") for r in block["runs"])
            # The reference export carries no "PlantUML source: <file>" / "SVG source: <file>" caption
            # under its diagrams, so drop them here. Matched on the caption's own wording rather than on
            # the italic flag alone, so genuine italic prose in the Supplement still comes through. The
            # shared .md and the PDF keep the captions -- only this format omits them.
            if block.get("italic") and _SOURCE_CAPTION.match(text):
                continue
            p = doc.add_paragraph()
            _add_runs(p, block["runs"])
            if block.get("italic"):
                for run in p.runs:
                    run.italic = True
        elif kind == "list":
            for item in block["items"]:
                style = "List Bullet" if item["level"] == 0 else "List Bullet 2"
                try:
                    p = doc.add_paragraph(style=style)
                except KeyError:
                    p = doc.add_paragraph()
                _add_runs(p, item["runs"])
        elif kind == "table":
            _add_table(doc, block)
            doc.add_paragraph()
        elif kind == "image":
            path = block["path"]
            if path and os.path.exists(path):
                doc.add_picture(path, width=Inches(IMAGE_WIDTH_IN))
            else:
                skipped_images.append(path)

    doc.save(out_path)
    return out_path, skipped_images


if __name__ == "__main__":
    md = sys.argv[1]
    tmpl = sys.argv[2] if len(sys.argv) > 2 else None
    out = sys.argv[3] if len(sys.argv) > 3 else None
    path, skipped = build_docx_scroll(md, tmpl, out)
    print("wrote", path)
    for s in skipped:
        print("  MISSING IMAGE:", s)
